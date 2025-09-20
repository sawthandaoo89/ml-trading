#!/usr/bin/env python3
"""
Convert markdown report to DOCX format
"""

import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def convert_markdown_to_docx(md_file, docx_file):
    """Convert markdown file to DOCX format"""
    
    # Read the markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create a new Document
    doc = Document()
    
    # Set up styles
    title_style = doc.styles['Title']
    title_style.font.size = Inches(0.2)
    
    heading1_style = doc.styles['Heading 1']
    heading1_style.font.size = Inches(0.16)
    
    heading2_style = doc.styles['Heading 2']
    heading2_style.font.size = Inches(0.14)
    
    # Split content into lines
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        # Title
        if line.startswith('# ') and i == 0:
            title = line[2:].strip()
            doc.add_heading(title, 0)
            
        # Heading 1
        elif line.startswith('## '):
            heading = line[3:].strip()
            doc.add_heading(heading, 1)
            
        # Heading 2
        elif line.startswith('### '):
            heading = line[4:].strip()
            doc.add_heading(heading, 2)
            
        # Bold text
        elif line.startswith('**') and line.endswith('**'):
            text = line[2:-2].strip()
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
            
        # Table
        elif '|' in line and not line.startswith('|'):
            # This is a table header
            table_lines = [line]
            i += 1
            # Skip separator line
            if i < len(lines) and '|' in lines[i] and '---' in lines[i]:
                i += 1
            # Collect table rows
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i])
                i += 1
            i -= 1  # Adjust for the loop increment
            
            # Create table
            if len(table_lines) > 0:
                # Parse table
                rows = []
                for table_line in table_lines:
                    if '|' in table_line:
                        cells = [cell.strip() for cell in table_line.split('|')[1:-1]]
                        rows.append(cells)
                
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Table Grid'
                    
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < len(table.rows[row_idx].cells):
                                table.rows[row_idx].cells[col_idx].text = cell_data
                                
        # Regular paragraph
        else:
            # Handle inline formatting
            paragraph_text = line
            
            # Handle bold text in paragraphs
            paragraph_text = re.sub(r'\*\*(.*?)\*\*', r'\1', paragraph_text)
            
            # Handle italic text
            paragraph_text = re.sub(r'\*(.*?)\*', r'\1', paragraph_text)
            
            # Handle code blocks
            if paragraph_text.startswith('```'):
                # Skip code block start
                i += 1
                code_lines = []
                while i < len(lines) and not lines[i].strip().startswith('```'):
                    code_lines.append(lines[i])
                    i += 1
                # Add code as a paragraph with monospace font
                if code_lines:
                    p = doc.add_paragraph()
                    run = p.add_run('\n'.join(code_lines))
                    run.font.name = 'Courier New'
            else:
                # Regular paragraph
                if paragraph_text:
                    p = doc.add_paragraph(paragraph_text)
                    
                    # Make bold text bold
                    if '**' in line:
                        # Find bold text and format it
                        bold_pattern = r'\*\*(.*?)\*\*'
                        matches = re.finditer(bold_pattern, line)
                        p.clear()
                        
                        last_end = 0
                        for match in matches:
                            # Add text before bold
                            if match.start() > last_end:
                                p.add_run(line[last_end:match.start()])
                            # Add bold text
                            bold_run = p.add_run(match.group(1))
                            bold_run.bold = True
                            last_end = match.end()
                        
                        # Add remaining text
                        if last_end < len(line):
                            p.add_run(line[last_end:])
        
        i += 1
    
    # Add the chart image if it exists
    try:
        doc.add_paragraph()  # Add some space
        doc.add_paragraph("Supporting Charts:")
        doc.add_picture('p3_report_charts.png', width=Inches(6))
    except:
        pass  # Chart not found, continue without it
    
    # Save the document
    doc.save(docx_file)
    print(f"Report converted to {docx_file}")

if __name__ == "__main__":
    convert_markdown_to_docx('p3_report.md', 'p3_report.docx')
